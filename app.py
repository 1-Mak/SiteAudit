from flask import Flask, render_template, request, redirect, url_for, jsonify, session
from werkzeug.utils import secure_filename
import os
from docx import Document
import PyPDF2
import mammoth
import time
import threading


#новые импорты:
from werkzeug.utils import secure_filename
from docx import Document
import markdown

import os
import markdown
from docx import Document
from PyPDF2 import PdfReader
from bs4 import BeautifulSoup

import datetime


from flask import send_from_directory, url_for


#Ванины импорты 2
from transformers import (
    T5Tokenizer, T5ForConditionalGeneration,
    DataCollatorForSeq2Seq, Trainer, TrainingArguments,
    get_linear_schedule_with_warmup, TrainerCallback,
    AutoTokenizer, AutoModelForSeq2SeqLM
)
from datasets import load_dataset, DatasetDict

import torch
from torch.optim import AdamW
import matplotlib.pyplot as plt

import os
import math
import time
# from IPython.display import display, clear_output
from transformers import pipeline
import pathlib, re, json, random, textwrap
from math import isnan

from typing import Union, Optional, Dict
from pathlib import Path
from pathlib import Path
# from IPython.display import HTML, display
from IPython.display import Markdown
import html
import html, textwrap, markdown
# from IPython.display import display, Markdown
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
import torch, json, textwrap
from huggingface_hub import login

# Ванины импорты
from transformers import LongT5ForConditionalGeneration, T5Tokenizer, set_seed, LongT5Config, AutoTokenizer, AutoModelForSeq2SeqLM
import numpy as np
import torch
import random
import gc

from flask import render_template
import markdown2
from markupsafe import Markup





    


app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'

# Это написал Иван Макаров
app.secret_key = 'KonstantaSosat'


processing_status = {'ready': False}
processing_result = None
raw_md = None
readme_html = None


if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])


def convert_file_to_md(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.md':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()  # plain text ≈ markdown-compatible

    elif ext == '.html':
        with open(file_path, 'r', encoding='utf-8') as f:
            html = f.read()
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text(separator="\n")  # чистый текст с переносами

    elif ext == '.pdf':
        reader = PdfReader(file_path)
        text = ''
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()

    elif ext == '.docx':
        doc = Document(file_path)
        md = ""
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name.lower()
            if "heading 1" in style:
                md += f"# {text}\n\n"
            elif "heading 2" in style:
                md += f"## {text}\n\n"
            elif "heading 3" in style:
                md += f"### {text}\n\n"
            elif "list" in style:
                md += f"- {text}\n"
            else:
                md += f"{text}\n\n"
        return md.strip()

    else:
        raise ValueError("Неподдерживаемый формат файла")    


@app.route('/')
def welcome():
    # Показываем шаблон index.html — главную страницу
    return render_template('1.welcome.html')


@app.route('/aiaudit')
def aiaudit():
    # Показываем шаблон about.html
    return render_template('2.ai-audit.html')

@app.route('/aiaudit_start')
def aiaudit_start():
    global raw_md

    filename = session.get('uploaded_file')
    if not filename:
        processing_result = 'Файл не был загружен.', 400

        processing_status['ready'] = True 

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file_extension = os.path.splitext(filename)[1].lower()

    if file_extension not in ALLOWED_EXTENSIONS:
        processing_result = 'Неподдерживаемый формат файла.', 400

        processing_status['ready'] = True 

    #ТУТ она уже .md ЙОУ
    raw_md = convert_file_to_md(file_path)

    # Запускаем обработку в фоновом потоке
    thread = threading.Thread(target=aiaudit_process)
    thread.start()

    return render_template('3.waiting.html')

#Вместо нейронки /ИИ

@app.route('/aiaudit_result')
def aiaudit_result():
    global processing_status, processing_result, readme_html
    
    # Проверяем готовность обработки
    if not processing_status.get('ready', False):
        return redirect('/aiaudit_start')
    
    criteria_section = processing_result

    # Сбрасываем статус обработки
    processing_status = {'ready': False}
    processing_result = None

    
    filename = session.get('uploaded_file')
    if not filename:
        return 'Файл не был загружен.', 400

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file_extension = os.path.splitext(filename)[1].lower()

    if file_extension not in ALLOWED_EXTENSIONS:
        return 'Неподдерживаемый формат файла.', 400

    
    if file_extension == '.md':
        content = extract_md_to_html(file_path)
    elif file_extension == '.docx':
        content = extract_docx_to_html(file_path)
    elif file_extension == '.pdf':
        content = extract_pdf_to_html(file_path)
    elif file_extension == '.html':
        content = extract_html_file(file_path)
    elif file_extension == '.txt':
        content = extract_txt_to_html(file_path)
    else:
        return 'Неизвестный формат.', 400

    criteria_html = markdown.markdown(criteria_section, extensions=["fenced_code", "tables"])
    readme_html = markdown.markdown(content, extensions=["fenced_code", "tables"])
    
    return render_template('4.result.html',
                         criteria_html=criteria_html,
                         readme_html = readme_html)

def aiaudit_process():
    global processing_status, processing_result, raw_md, readme_html
    # Тут твоя нейросеть работает (замени на реальную функцию)

    gc.collect()

    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

    def seed_all(seed: int):
        random.seed(seed)
        np.random.seed(seed)
        torch.manual_seed(seed)
        torch.cuda.manual_seed_all(seed)
        set_seed(seed)
        
    seed = 52
    seed_all(seed)

    config = LongT5Config.from_pretrained("C:/Users/MSI GL66/Documents/GitHub/Site/model")

    tokenizer = AutoTokenizer.from_pretrained("C:/Users/MSI GL66/Documents/GitHub/Site/model")
    model = AutoModelForSeq2SeqLM.from_pretrained("C:/Users/MSI GL66/Documents/GitHub/Site/model")

    model.to(device)


    def improve_readme_from_text(bad_md_text: str,
                                 device: str = "cpu",
                                 gen_kwargs: dict | None = None) -> str:
        """
        Принимает markdown-текст, заменяет переводы строк на <NL>,
        генерирует «улучшенный» вариант через LongT5, возвращает markdown.
        """
        # 1) Добавляем <NL> как дополнительный спец-токен.
        #    tokenizer.add_special_tokens возвращает число реально добавленных токенов.
        num_added = tokenizer.add_special_tokens({
            "additional_special_tokens": ["<NL>"]
        })
        if num_added > 0:
            # если что-то добавилось — расширяем эмбеддинги
            model.resize_token_embeddings(len(tokenizer))

        # 2) Готовим вход: заменяем \n → <NL>
        bad_text = bad_md_text.replace("\n", "<NL>")
        inputs = tokenizer(
            "improve_readme: " + bad_text,
            return_tensors="pt",
            max_length=4096,
            truncation=True
        ).to(device)

        # 3) Параметры генерации по умолчанию
        if gen_kwargs is None:
            gen_kwargs = {
                "max_new_tokens": 2048,
                "num_beams": 4,
                "early_stopping": True,
            }

        # 4) Генерация
        generated_ids = model.generate(**inputs, **gen_kwargs)
        output_text = tokenizer.decode(
            generated_ids[0],
            skip_special_tokens=True,
            clean_up_tokenization_spaces=False
        )

        # 5) Обратная замена <NL> → \n
        return output_text.replace("<NL>", "\n")


    

    fixed_md = improve_readme_from_text(raw_md, device="cuda")

    #ВОТ ТУТ ВТОРАЯ НЕЙРОНКА ОБРАБАТЫВАЕТ ЗАПРОС


    #Конец данных от нейронки

    

    #ВОТ ТУТ КОНЕЦ ВТОРОЙ НЕЙРОНКИ
    #ВОТ ТУТ ДАННЫЕ ПОМЕНЯТЬ НА ВЫВОД ВМЕСТО fixed_md на criteria_section
    processing_result = fixed_md
    readme_html = """ My Project

## About My Project

This project is a solution to the problem of [describe the problem]. It is designed for [describe the audience] and differentiates itself from other solutions by [describe the differentiator].

## Installation

1. Install prerequisites: [list prerequisites].
2. Clone the repository: `git clone https://github.com/username/repo.git`.
3. Install dependencies: `pip install -r requirements.txt`.

## Usage

Use the following command to run the project: `python my_project.py`.

## Contributing

Contributions are welcome! Please read our [Contributing Guidelines](CONTRIBUTING.md) for details.

## License

This project is licensed under the [MIT License](LICENSE).

## Support

For support, please join our [Discord community](https://discord.gg/community)."""
    # processing_result = criteria_section

    processing_status['ready'] = True


@app.route('/check_aiaudit_status')
def check_aiaudit_status():
    return jsonify({'ready': processing_status['ready']})




@app.route('/manaudit')
def manaudit():
    # Показываем шаблон about.html
    return render_template('5.manual-audit.html')

def extract_md_to_html(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
        return markdown.markdown(md_content, extensions=["fenced_code", "tables"])


def extract_docx_to_html(file_path):
    doc = Document(file_path)
    html = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower()
        if "heading 1" in style:
            html += f"<h1>{text}</h1>"
        elif "heading 2" in style:
            html += f"<h2>{text}</h2>"
        elif "heading 3" in style:
            html += f"<h3>{text}</h3>"
        else:
            html += f"<p>{text}</p>"
    return html


from PyPDF2 import PdfReader

def extract_pdf_to_html(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    from html import escape
    return f"<pre style='white-space: pre-wrap; font-family: monospace;'>{escape(text.strip())}</pre>"



def extract_html_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()  # предполагается, что HTML уже готов к вставке


def extract_txt_to_html(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
        return f"<pre>{text}</pre>"


ALLOWED_EXTENSIONS = {'.md', '.pdf', '.html', '.docx', '.txt'}

@app.route('/manaudit_check')
def manaudit_check():
    filename = session.get('uploaded_file')
    if not filename:
        return 'Файл не был загружен.', 400

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file_extension = os.path.splitext(filename)[1].lower()

    if file_extension not in ALLOWED_EXTENSIONS:
        return 'Неподдерживаемый формат файла.', 400

    try:
        if file_extension == '.md':
            content = extract_md_to_html(file_path)
        elif file_extension == '.docx':
            content = extract_docx_to_html(file_path)
        elif file_extension == '.pdf':
            content = extract_pdf_to_html(file_path)
        elif file_extension == '.html':
            content = extract_html_file(file_path)
        elif file_extension == '.txt':
            content = extract_txt_to_html(file_path)
        else:
            return 'Неизвестный формат.', 400

        return render_template('6.manual-check.html', content=content)

    except Exception as e:
        return f'Ошибка при обработке файла: {str(e)}', 500



@app.route('/analyze', methods=['POST'])
def analyze():
    if 'file' not in request.files:
        return 'Файл не найден в запросе.', 400

    file = request.files['file']
    if file.filename == '':
        return 'Имя файла отсутствует.', 400

    filename = secure_filename(file.filename)
    file_extension = os.path.splitext(filename)[1].lower()

    if file_extension not in ALLOWED_EXTENSIONS:
        return 'Формат файла не поддерживается.', 400

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    session['uploaded_file'] = filename
    return '', 204  # Возвращаем "No Content", всё успешно

@app.route('/save_audit', methods=['POST'])
def save_audit():
    data = request.get_json()
    criteria = data.get('criteria', [])
    general_comment = data.get('general_comment', '')

    try:
        uploaded_filename = session.get('uploaded_file', 'Unknown file')

        # Генерация HTML-контента
        html_content = "<html><head><meta charset='utf-8'><title>Manual Audit Result</title></head><body>"
        html_content += f"<h1>Manual Audit Result</h1>"
        html_content += f"<p><strong>Audited File:</strong> {uploaded_filename}</p><ol>"

        for item in criteria:
            score = item.get("grade", 0)
            comment = item.get("comment", "").strip()
            html_content += f"<li>Criterion {item['id']} – score: {score} – comment: {comment}</li>"

        html_content += "</ol>"

        if general_comment:
            html_content += f"<p><strong>Recommendations for improvement:</strong> {general_comment}</p>"

        html_content += "</body></html>"

        filename = f"manual_audit_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        file_path = os.path.join("saved_audits", filename)

        os.makedirs("saved_audits", exist_ok=True)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        download_url = url_for('download_file', filename=filename)
        return jsonify({'status': 'success', 'file': filename, 'download_url': download_url})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/view/<filename>')
def view_file(filename):
    # Формируем путь к файлу
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    # Проверяем, существует ли файл
    if not os.path.exists(file_path):
        return 'Файл не найден', 404

    # Извлекаем текст из .docx файла
    try:
        doc = Document(file_path)
        content = '\n'.join([para.text for para in doc.paragraphs])
        return content
    except Exception as e:
        return f'Ошибка при обработке файла: {str(e)}', 500
    
@app.route('/save_audit_ai', methods=['POST'])
def save_audit_ai():
    data = request.get_json()
    criteria_html = data.get('criteria_html', '')

    try:
        uploaded_filename = session.get('uploaded_file', 'Unknown file')

        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>AI Audit Result</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            padding: 20px;
        }}
        .markdown-body {{
            line-height: 1.6;
        }}
    </style>
</head>
<body>
    <h1>AI Audit Result</h1>
    <p><strong>Audited File:</strong> {uploaded_filename}</p>
    <div class="markdown-body">
        {criteria_html}
    </div>
</body>
</html>"""

        filename = f"ai_audit_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        file_path = os.path.join("saved_audits", filename)

        os.makedirs("saved_audits", exist_ok=True)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        download_url = url_for('download_file', filename=filename)
        return jsonify({'status': 'success', 'file': filename, 'download_url': download_url})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory('saved_audits', filename, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
